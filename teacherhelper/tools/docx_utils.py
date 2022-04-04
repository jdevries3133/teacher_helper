import re
from typing import List, Union

import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def docx_regex_search(
    doc: docx.Document,
    regex: Union[str, re.Pattern],
) -> List[str]:
    """Iterates through all paragraphs, tables, and hyperlinks in a word doc,
    and extracts any strings that match the regex.

    TODO: this needs tests, but it works nicely as far as I recall.
    """
    matches: List[str] = []
    document = docx.Document(doc)

    # search paragraphs
    for par in document.paragraphs:
        mo = re.fullmatch(regex, par.text.strip())
        if mo:
            matches.append(mo.string)

    # search tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                mo = re.fullmatch(regex, cell.text.strip())
                if mo:
                    matches.append(mo.string)

    # search for links
    for rel in document.part.rels:
        Rel = document.part.rels[rel]
        if Rel.reltype == RT.HYPERLINK:
            text = Rel._target
            mo = re.fullmatch(regex, text)
            if mo:
                matches.append(mo.string)

    return matches
