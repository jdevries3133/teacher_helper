import re

import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def regex_search(st, doc, regex, near_match_regex=r''):
    matches = []
    misses = []
    document = docx.Document(doc)
    # search paragraphs
    for par in document.paragraphs:
        mo = re.fullmatch(regex, par.text.strip())
        if mo:
            matches.append([st.name, mo.string, st.homeroom])
        else:
            mo = re.search(near_match_regex, par.text)
            if mo:
                misses.append([st.name, mo.string, st.homeroom])
    # search tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                mo = re.fullmatch(regex, cell.text.strip())
                if mo:
                    matches.append([st.name, mo.string, st.homeroom])
                else:
                    mo = re.search(near_match_regex, cell.text)
                    if mo:
                        misses.append([st.name, mo.string, st.homeroom])
    # search for links
    for rel in document.part.rels:
        Rel = document.part.rels[rel]
        if Rel.reltype == RT.HYPERLINK:
            text = Rel._target
            mo = re.fullmatch(regex, cell.text.strip())
            if mo:
                matches.append([st.name, mo.string, st.homeroom])
            else:
                mo = re.search(near_match_regex, cell.text)
                if mo:
                    misses.append([st.name, mo.string, st.homeroom])
    if near_match_regex:
        return matches, misses
    return matches, None